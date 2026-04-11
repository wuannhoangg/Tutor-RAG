from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Tuple
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from ..schemas.document import DocumentMetadata


def _get_attr(obj: Any, *names: str, default: Any = None) -> Any:
    if obj is None:
        return default

    if isinstance(obj, dict):
        for name in names:
            if name in obj:
                return obj[name]
        return default

    for name in names:
        if hasattr(obj, name):
            return getattr(obj, name)
    return default


class DOCXParser:
    """
    Extract text from DOCX files.
    Uses python-docx when available and falls back to XML parsing.
    """

    def parse(self, file_path: str, source_metadata: DocumentMetadata) -> Tuple[str, List[dict]]:
        path = Path(file_path)

        if path.suffix.lower() != ".docx":
            raise ValueError("Invalid file type for DOCXParser.")
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            paragraphs = self._parse_with_python_docx(path)
        except Exception:
            paragraphs = self._parse_with_zip_xml(path)

        cleaned_paragraphs = [p.strip() for p in paragraphs if p and p.strip()]
        raw_text = "\n\n".join(cleaned_paragraphs).strip()

        page_details: List[Dict[str, Any]] = []
        if raw_text:
            page_details.append(
                {
                    "page_number": 1,
                    "start_char": 0,
                    "end_char": len(raw_text),
                    "text_excerpt": raw_text[:300],
                }
            )

        return raw_text, page_details

    def _parse_with_python_docx(self, path: Path) -> List[str]:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        paragraphs: List[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return paragraphs

    def _parse_with_zip_xml(self, path: Path) -> List[str]:
        paragraphs: List[str] = []
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        with ZipFile(path, "r") as archive:
            xml_bytes = archive.read("word/document.xml")

        root = ET.fromstring(xml_bytes)

        for para in root.findall(".//w:p", namespaces):
            texts = []
            for node in para.findall(".//w:t", namespaces):
                if node.text:
                    texts.append(node.text)
            paragraph_text = "".join(texts).strip()
            if paragraph_text:
                paragraphs.append(paragraph_text)

        return paragraphs